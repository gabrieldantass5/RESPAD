import os
import json
import asyncio
from pathlib import Path
from dotenv import load_dotenv
from fastapi import FastAPI, Request
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from sse_starlette.sse import EventSourceResponse
from google import genai
from google.genai import types
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
from system_prompt import SYSTEM_PROMPT

load_dotenv()

client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))
MODEL = "gemini-2.5-flash"

app = FastAPI(title="Assistente RESPAD/CBMAL")
app.mount("/assets", StaticFiles(directory="assets"), name="assets")

@app.get("/", response_class=HTMLResponse)
async def root():
    return FileResponse("static/index.html")

@app.post("/chat")
async def chat(request: Request):
    body = await request.json()
    messages = body.get("messages", [])

    # Converter histórico para formato google-genai
    history = []
    for msg in messages[:-1]:
        role = "user" if msg["role"] == "user" else "model"
        history.append(types.Content(role=role, parts=[types.Part(text=msg["content"])]))

    user_message = messages[-1]["content"] if messages else ""

    async def generate():
        try:
            def stream_sync():
                chat_session = client.chats.create(
                    model=MODEL,
                    config=types.GenerateContentConfig(
                        system_instruction=SYSTEM_PROMPT,
                        temperature=0.7,
                    ),
                    history=history,
                )
                return list(chat_session.send_message_stream(user_message))

            chunks = await asyncio.to_thread(stream_sync)
            for chunk in chunks:
                if chunk.text:
                    yield {"data": json.dumps({"text": chunk.text})}
        except Exception as e:
            yield {"data": json.dumps({"text": f"[ERRO] {str(e)}"})}
        yield {"data": json.dumps({"done": True})}

    return EventSourceResponse(generate())

@app.post("/export")
async def export_doc(request: Request):
    body = await request.json()
    messages = body.get("messages", [])
    title = body.get("title", "Documento RESPAD")

    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(1.18)
    section.top_margin = Inches(0.98)
    section.bottom_margin = Inches(0.98)

    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run()
    brasao_path = Path("assets/brasao.png")
    if brasao_path.exists():
        run.add_picture(str(brasao_path), width=Inches(0.6))

    header_para2 = header.add_paragraph()
    header_para2.paragraph_format.space_before = Pt(4)
    border_run = header_para2.add_run("─" * 80)
    border_run.font.color.rgb = RGBColor(0xC1, 0x0A, 0x0A)
    border_run.font.size = Pt(8)

    title_para = doc.add_heading(title, level=1)
    title_para.runs[0].font.color.rgb = RGBColor(0xC1, 0x0A, 0x0A)
    title_para.runs[0].font.name = "Calibri"

    doc.add_paragraph()

    for msg in messages:
        if msg["role"] == "assistant":
            para = doc.add_paragraph(msg["content"])
            para.runs[0].font.size = Pt(11)
            para.runs[0].font.name = "Calibri"
            doc.add_paragraph()

    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        "Gerado pelo Assistente RESPAD/CBMAL · Corpo de Bombeiros Militar de Alagoas"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name

    safe_title = "".join(c for c in title if c.isalnum() or c in " -_").strip()
    filename = f"{safe_title}.docx" if safe_title else "documento-respad.docx"

    return FileResponse(
        tmp_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )
